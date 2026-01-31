from __future__ import annotations

import base64
import logging
import os
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass
from pathlib import Path
from typing import Callable

import pandas as pd
from docx import Document
from dotenv import load_dotenv
from openai import OpenAI
from PyPDF2 import PdfReader

load_dotenv()
logger = logging.getLogger("rag")

SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".csv", ".xlsx"}
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp", ".tiff", ".tif"}
SUPPORTED_EXTENSIONS = SUPPORTED_EXTENSIONS | IMAGE_EXTENSIONS

# Minimum image dimensions to process (skip icons/thumbnails)
MIN_IMAGE_WIDTH = 200
MIN_IMAGE_HEIGHT = 200


@dataclass(frozen=True)
class IngestedDocument:
    filename: str
    filepath: str
    content: str


class DocumentIngester:
    def __init__(self, docs_dir: Path, local_only: bool = False) -> None:
        self.docs_dir = docs_dir
        self.local_only = local_only  # If True, never use cloud APIs
        self._openai_client: OpenAI | None = None

    def _get_openai_client(self) -> OpenAI | None:
        """
        Get or create OpenAI client for image processing.
        
        Tries in order:
        1. macOS Keychain (via security module)
        2. Environment variable
        3. .env file
        """
        if self.local_only:
            return None
        
        if self._openai_client is None:
            api_key = None
            
            # Try keychain first
            try:
                from app.config import DATA_DIR
                from app.security import get_key_manager
                key_manager = get_key_manager(DATA_DIR)
                api_key = key_manager.get_api_key("OPENAI_API_KEY")
            except Exception:
                pass
            
            # Fall back to environment
            if not api_key:
                api_key = os.getenv("OPENAI_API_KEY")
            
            if api_key:
                self._openai_client = OpenAI(api_key=api_key.strip())
        
        return self._openai_client
    
    def _is_image_large_enough(self, path: Path) -> bool:
        """
        Check if image meets minimum dimension requirements.
        Skips small icons, thumbnails, and UI elements.
        """
        try:
            from PIL import Image
            with Image.open(path) as img:
                width, height = img.size
                return width >= MIN_IMAGE_WIDTH and height >= MIN_IMAGE_HEIGHT
        except ImportError:
            # PIL not installed, use file size heuristic
            try:
                size = path.stat().st_size
                return size > 50 * 1024  # > 50KB likely not an icon
            except OSError:
                return False
        except Exception:
            return False

    def _read_image(self, path: Path) -> str:
        """
        Process an image using OpenAI's vision API.
        
        Returns a text description including:
        - Detailed description of the image content
        - Any visible text (OCR)
        - Key visual elements useful for search
        
        Skips small images (icons, thumbnails) to save API costs.
        """
        # Skip if local-only mode
        if self.local_only:
            logger.debug("Skipping image (local-only mode): %s", path.name)
            return ""
        
        # Check image dimensions first (skip small icons)
        if not self._is_image_large_enough(path):
            logger.debug("Skipping small image (icon/thumbnail): %s", path.name)
            return ""
        
        client = self._get_openai_client()
        if not client:
            logger.warning("OpenAI API key not found, skipping image: %s", path.name)
            return ""

        try:
            # Read and encode image as base64
            with open(path, "rb") as f:
                image_data = base64.b64encode(f.read()).decode("utf-8")

            # Determine MIME type
            suffix = path.suffix.lower()
            mime_map = {
                ".jpg": "image/jpeg",
                ".jpeg": "image/jpeg",
                ".png": "image/png",
                ".gif": "image/gif",
                ".webp": "image/webp",
                ".bmp": "image/bmp",
                ".tiff": "image/tiff",
                ".tif": "image/tiff",
            }
            mime_type = mime_map.get(suffix, "image/jpeg")

            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": """Analyze this image and provide a detailed description for a personal knowledge base.

Include:
1. A comprehensive description of what's in the image (people, objects, scenes, locations)
2. ALL visible text - transcribe every piece of text you can see (signs, documents, labels, etc.)
3. Key details like dates, names, numbers, or any factual information visible
4. Context clues about what this image represents (document type, event, location, etc.)

Format your response as natural prose that will be searchable. Be thorough but concise.""",
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:{mime_type};base64,{image_data}",
                                    "detail": "high",
                                },
                            },
                        ],
                    }
                ],
                max_tokens=1000,
            )

            description = response.choices[0].message.content.strip()
            # Prefix with image metadata for context
            return f"[Image: {path.name}]\n{description}"

        except Exception as e:
            logger.warning("Failed to process image %s: %s", path.name, e)
            return ""

    def ingest_all(self, parallel: bool = False, max_workers: int = 4) -> list[dict[str, str]]:
        """
        Ingest all documents from the docs directory.
        
        Args:
            parallel: If True, process files in parallel (faster for I/O bound operations)
            max_workers: Number of parallel workers (default: 4)
        
        Returns:
            List of dicts with filename, filepath, and content.
        """
        if not self.docs_dir.exists():
            raise RuntimeError(
                f"Docs directory not found: {self.docs_dir}. "
                "Create it and add files before indexing."
            )

        files = [
            path
            for path in self.docs_dir.rglob("*")
            if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
        ]
        if not files:
            raise RuntimeError(
                f"No supported documents found in {self.docs_dir}. "
                f"Supported extensions: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            )

        if parallel:
            documents = self._ingest_parallel(files, max_workers)
        else:
            documents = self._ingest_sequential(files)

        if not documents:
            raise RuntimeError("All documents were empty after ingestion.")

        return [
            {
                "filename": doc.filename,
                "filepath": doc.filepath,
                "content": doc.content,
            }
            for doc in documents
        ]
    
    def _ingest_sequential(self, files: list[Path]) -> list[IngestedDocument]:
        """Process files sequentially."""
        documents: list[IngestedDocument] = []
        for path in files:
            content = self._read_file(path)
            if content.strip():
                documents.append(
                    IngestedDocument(
                        filename=path.name,
                        filepath=str(path),
                        content=content,
                    )
                )
        return documents
    
    def _ingest_parallel(
        self, 
        files: list[Path], 
        max_workers: int
    ) -> list[IngestedDocument]:
        """
        Process files in parallel using ThreadPoolExecutor.
        
        Uses threads (not processes) because file I/O is the bottleneck.
        """
        documents: list[IngestedDocument] = []
        
        def process_file(path: Path) -> IngestedDocument | None:
            try:
                content = self._read_file(path)
                if content.strip():
                    return IngestedDocument(
                        filename=path.name,
                        filepath=str(path),
                        content=content,
                    )
            except Exception as e:
                logger.warning("Failed to process %s: %s", path.name, e)
            return None
        
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Submit all files
            future_to_path = {
                executor.submit(process_file, path): path 
                for path in files
            }
            
            # Collect results as they complete
            for future in as_completed(future_to_path):
                result = future.result()
                if result:
                    documents.append(result)
        
        return documents
    
    def ingest_files(
        self,
        file_paths: list[Path],
        parallel: bool = True,
        max_workers: int = 4,
        progress_callback: Callable[[str, int, int], None] | None = None,
    ) -> list[dict[str, str]]:
        """
        Ingest specific files (for use with device scanner).
        
        Args:
            file_paths: List of file paths to ingest
            parallel: Use parallel processing
            max_workers: Number of parallel workers
            progress_callback: Optional callback(filename, current, total)
        
        Returns:
            List of dicts with filename, filepath, and content.
        """
        if not file_paths:
            return []
        
        documents: list[IngestedDocument] = []
        total = len(file_paths)
        
        def process_file(path: Path, index: int) -> IngestedDocument | None:
            try:
                if progress_callback:
                    progress_callback(path.name, index, total)
                
                content = self._read_file(path)
                if content.strip():
                    return IngestedDocument(
                        filename=path.name,
                        filepath=str(path),
                        content=content,
                    )
            except Exception as e:
                logger.warning("Failed to process %s: %s", path.name, e)
            return None
        
        if parallel and len(file_paths) > 1:
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                futures = {
                    executor.submit(process_file, path, i): path
                    for i, path in enumerate(file_paths)
                }
                for future in as_completed(futures):
                    result = future.result()
                    if result:
                        documents.append(result)
        else:
            for i, path in enumerate(file_paths):
                result = process_file(path, i)
                if result:
                    documents.append(result)
        
        return [
            {
                "filename": doc.filename,
                "filepath": doc.filepath,
                "content": doc.content,
            }
            for doc in documents
        ]

    def _read_file(self, path: Path) -> str:
        suffix = path.suffix.lower()
        try:
            if suffix in {".txt", ".md"}:
                return path.read_text(encoding="utf-8", errors="ignore")
            if suffix == ".pdf":
                reader = PdfReader(str(path))
                return "\n".join(page.extract_text() or "" for page in reader.pages)
            if suffix == ".docx":
                doc = Document(str(path))
                return "\n".join(p.text for p in doc.paragraphs)
            if suffix == ".csv":
                df = pd.read_csv(path)
                return df.to_csv(index=False)
            if suffix == ".xlsx":
                df = pd.read_excel(path)
                return df.to_csv(index=False)
            if suffix in IMAGE_EXTENSIONS:
                return self._read_image(path)
        except Exception as exc:
            raise RuntimeError(f"Failed to read {path}") from exc

        raise RuntimeError(f"Unsupported file type: {path}")
